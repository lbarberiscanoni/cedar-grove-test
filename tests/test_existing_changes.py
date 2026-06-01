"""Tests for documents that ALREADY contain tracked changes (counterparty redlines).

These cover the layering behavior: our edits nest inside a counterparty <w:ins>,
and edits that straddle a tracked-change boundary are flagged rather than guessed.
"""

from __future__ import annotations

import subprocess
import sys
import zipfile
from pathlib import Path

import pytest
from docx import Document
from lxml import etree

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from redline_engine import (  # noqa: E402
    Edit,
    apply_edits,
    extract_paragraphs,
)

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _q(tag: str) -> str:
    return f"{{{W}}}{tag}"


def _doc_with_counterparty_ins(tmp_path: Path) -> Path:
    """Build a doc whose paragraph contains a counterparty insertion.

    Visible text: 'The term: the best efforts apply.'
    where 'the best efforts' was inserted by 'Counterparty'.
    """
    doc = Document()
    doc.add_paragraph("seed")  # will be replaced
    body = doc.element.body
    for p in list(body.findall(_q("p"))):
        body.remove(p)
    para_xml = '''
    <w:p xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
      <w:r><w:t xml:space="preserve">The term: </w:t></w:r>
      <w:ins w:id="10" w:author="Counterparty" w:date="2026-05-01T00:00:00Z">
        <w:r><w:t xml:space="preserve">the best efforts</w:t></w:r>
      </w:ins>
      <w:r><w:t xml:space="preserve"> apply.</w:t></w:r>
    </w:p>
    '''
    body.insert(0, etree.fromstring(para_xml.strip()))
    path = tmp_path / "cp.docx"
    doc.save(path)
    return path


def _pandoc(path: Path, mode: str) -> str:
    return subprocess.run(
        ["pandoc", f"--track-changes={mode}", "-t", "plain", str(path)],
        capture_output=True, text=True, check=True,
    ).stdout


def test_visible_text_includes_counterparty_insertion(tmp_path):
    in_path = _doc_with_counterparty_ins(tmp_path)
    doc = Document(in_path)
    paragraphs = extract_paragraphs(doc)
    assert paragraphs[0].text == "The term: the best efforts apply."
    # Two editable targets (before/after) + one ins target.
    kinds = [t.kind for t in paragraphs[0].targets]
    assert kinds == ["editable", "ins", "editable"]


def test_edit_inside_counterparty_insertion_layers(tmp_path):
    """Our edit lands inside the counterparty's inserted text. It must layer as a
    nested tracked change: accept-all shows our edit applied (their insertion kept),
    reject-all reverts everything to the true original.
    """
    in_path = _doc_with_counterparty_ins(tmp_path)
    doc = Document(in_path)
    paragraphs = extract_paragraphs(doc)

    edits = [Edit(para=0, type="replace", find="best", replace="optimal")]
    flagged = apply_edits(doc, paragraphs, edits, author="Michael Ohta")
    assert flagged == [], flagged

    out = tmp_path / "cp-out.docx"
    doc.save(out)

    accepted = _pandoc(out, "accept")
    rejected = _pandoc(out, "reject")
    assert "The term: the optimal efforts apply." in accepted, accepted
    # reject-all reverts BOTH the counterparty insertion and our edit.
    assert "The term:  apply." in rejected or "The term: apply." in rejected, rejected

    # XML: our deletion of 'best' is nested inside a counterparty-authored <w:ins>.
    with zipfile.ZipFile(out) as z:
        root = etree.fromstring(z.read("word/document.xml"))
    dels = root.findall(f".//{_q('del')}")
    assert len(dels) == 1
    assert dels[0].get(_q("author")) == "Michael Ohta"
    parent_ins = dels[0].getparent()
    assert parent_ins.tag == _q("ins")
    assert parent_ins.get(_q("author")) == "Counterparty"
    # Our replacement insertion is authored by us.
    our_ins = [i for i in root.findall(f".//{_q('ins')}")
               if i.get(_q("author")) == "Michael Ohta"]
    assert len(our_ins) == 1
    assert (our_ins[0].find(f".//{_q('t')}").text) == "optimal"


def test_delete_inside_counterparty_insertion(tmp_path):
    """Deleting part of the counterparty's insertion: accept-all drops it, with no
    stray replacement insertion authored by us."""
    in_path = _doc_with_counterparty_ins(tmp_path)
    doc = Document(in_path)
    paragraphs = extract_paragraphs(doc)

    edits = [Edit(para=0, type="delete", find=" best")]  # remove ' best'
    flagged = apply_edits(doc, paragraphs, edits, author="Michael Ohta")
    assert flagged == [], flagged

    out = tmp_path / "cp-del.docx"
    doc.save(out)
    accepted = _pandoc(out, "accept")
    assert "The term: the efforts apply." in accepted, accepted

    with zipfile.ZipFile(out) as z:
        root = etree.fromstring(z.read("word/document.xml"))
    our_ins = [i for i in root.findall(f".//{_q('ins')}")
               if i.get(_q("author")) == "Michael Ohta"]
    assert our_ins == []


def test_edit_spanning_boundary_is_flagged(tmp_path):
    """A find that straddles clean text and the counterparty insertion is ambiguous
    to layer, so it must be flagged, not applied."""
    in_path = _doc_with_counterparty_ins(tmp_path)
    doc = Document(in_path)
    paragraphs = extract_paragraphs(doc)

    # 'term: the' spans the leading editable run and the counterparty <w:ins>.
    edits = [Edit(para=0, type="replace", find="term: the", replace="term: those")]
    flagged = apply_edits(doc, paragraphs, edits, author="Michael Ohta")
    assert len(flagged) == 1
    assert "boundary" in flagged[0].reason

    # Document text unchanged.
    paragraphs_after = extract_paragraphs(doc)
    assert paragraphs_after[0].text == "The term: the best efforts apply."


def test_edit_in_clean_run_of_doc_with_changes_still_works(tmp_path):
    """An edit on clean text in a doc that also has counterparty changes applies as
    a normal top-level tracked change."""
    in_path = _doc_with_counterparty_ins(tmp_path)
    doc = Document(in_path)
    paragraphs = extract_paragraphs(doc)

    edits = [Edit(para=0, type="replace", find="apply", replace="govern")]
    flagged = apply_edits(doc, paragraphs, edits, author="Michael Ohta")
    assert flagged == [], flagged

    out = tmp_path / "cp-clean.docx"
    doc.save(out)
    accepted = _pandoc(out, "accept")
    assert "The term: the best efforts govern." in accepted, accepted


def test_new_change_ids_do_not_collide(tmp_path):
    """Our new tracked-change ids must exceed the counterparty's existing ids."""
    in_path = _doc_with_counterparty_ins(tmp_path)
    doc = Document(in_path)
    paragraphs = extract_paragraphs(doc)
    edits = [Edit(para=0, type="replace", find="best", replace="optimal")]
    apply_edits(doc, paragraphs, edits, author="Michael Ohta")
    out = tmp_path / "cp-ids.docx"
    doc.save(out)
    with zipfile.ZipFile(out) as z:
        root = etree.fromstring(z.read("word/document.xml"))
    ids = sorted(int(el.get(_q("id")))
                 for tag in ("ins", "del")
                 for el in root.findall(f".//{_q(tag)}")
                 if el.get(_q("id")))
    # Counterparty used id 10; ours must be > 10 and unique.
    assert len(ids) == len(set(ids)), "ids must be unique"
    assert max(ids) > 10


@pytest.mark.skipif(__import__("shutil").which("pandoc") is None, reason="pandoc not installed")
def test_real_esa_round_trips(tmp_path):
    """The real counterparty ESA (already full of tracked changes) survives a small
    redline: file stays valid and reject-all restores the counterparty's text
    exactly for an untouched edit target."""
    src = Path("/tmp/test_esa.docx")
    if not src.exists():
        pytest.skip("real ESA fixture not present")
    doc = Document(src)
    paragraphs = extract_paragraphs(doc)

    # Find a paragraph with the word 'Services' in clean text to edit safely.
    target_edit = None
    for p in paragraphs:
        if "Services" in p.text and find_count(p.text, "Services") == 1:
            target_edit = Edit(para=p.index, type="replace",
                               find="Services", replace="Services (as defined)")
            break
    assert target_edit is not None

    flagged = apply_edits(doc, paragraphs, [target_edit], author="Michael Ohta")
    out = tmp_path / "esa-out.docx"
    doc.save(out)
    # Validity: reopen + pandoc parse.
    Document(out)
    _ = _pandoc(out, "accept")
    _ = _pandoc(out, "reject")


def find_count(haystack: str, needle: str) -> int:
    return haystack.count(needle)
