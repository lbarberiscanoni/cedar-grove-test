"""Tests for the CLI / API layer in review.py.

We don't hit the real Claude API in unit tests — instead we stub the client so the
end-to-end review() flow is exercised deterministically.
"""

from __future__ import annotations

import json
import sys
from pathlib import Path
from types import SimpleNamespace

from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from review import (  # noqa: E402
    parse_edits,
    review,
    review_bytes,
    strip_fences,
)


def test_strip_fences_no_fence():
    s = '{"edits": []}'
    assert strip_fences(s) == s


def test_strip_fences_with_json_fence():
    s = '```json\n{"edits": []}\n```'
    assert strip_fences(s) == '{"edits": []}'


def test_strip_fences_with_bare_fence():
    s = '```\n{"edits": [1,2]}\n```'
    assert strip_fences(s) == '{"edits": [1,2]}'


def test_parse_edits_defaults():
    raw = json.dumps({"edits": [
        {"para": 3, "find": "x", "replace": "y"},
        {"para": 4, "type": "delete", "find": "z"},
    ]})
    edits = parse_edits(raw)
    assert len(edits) == 2
    assert edits[0].para == 3 and edits[0].type == "replace"
    assert edits[1].type == "delete" and edits[1].is_delete


class _StubClient:
    """Stub anthropic.Anthropic.messages.create() for end-to-end testing.

    Accepts a single response string (returned for every call) or a list of strings
    (returned in order — for testing the main + completeness passes separately).
    """

    def __init__(self, response_text):
        self._responses = response_text if isinstance(response_text, list) else None
        self._text = None if isinstance(response_text, list) else response_text
        self.calls: list[dict] = []
        self.messages = SimpleNamespace(create=self._create)

    def _create(self, **kwargs):
        i = len(self.calls)
        self.calls.append(kwargs)
        if self._responses is not None:
            text = self._responses[min(i, len(self._responses) - 1)]
        else:
            text = self._text
        return SimpleNamespace(content=[SimpleNamespace(text=text)])


def _response_two_edits() -> str:
    return json.dumps({"edits": [
        {"para": 0, "type": "replace",
         "find": "best efforts", "replace": "commercially reasonable efforts",
         "severity": "high", "comment": "Playbook §3."},
        {"para": 99, "type": "replace",
         "find": "anything", "replace": "something",
         "severity": "low", "comment": "out of range — should flag"},
    ]})


def test_review_end_to_end_with_stub(tmp_path):
    # Build a fixture contract.
    doc = Document()
    doc.add_paragraph("The term is best efforts to deliver.")
    in_path = tmp_path / "contract.docx"
    doc.save(in_path)

    # Both passes return the same 2 edits; the completeness pass's edits are deduped.
    client = _StubClient(_response_two_edits())

    out_docx, out_json, result = review(str(in_path), out_dir=str(tmp_path), client=client)
    assert out_docx.exists()
    assert out_json.exists()
    assert result.num_applied == 1
    assert result.num_flagged == 1
    assert "out of range" in result.flagged[0]["reason"]

    # Main + completeness pass = two calls, each with the playbook cached for 1h.
    assert len(client.calls) == 2
    for call in client.calls:
        assert any(b.get("cache_control", {}).get("ttl") == "1h"
                   for b in call["system"]), "playbook block should be cached 1h"

    # Output paths follow the documented naming convention.
    assert out_docx.name == "contract.redlined.docx"
    assert out_json.name == "contract.flagged.json"

    flagged_data = json.loads(out_json.read_text())
    assert flagged_data[0]["edit"]["para"] == 99


def test_completeness_pass_adds_missed_edits(tmp_path):
    """The second pass contributes edits the main pass missed; dups are dropped."""
    doc = Document()
    doc.add_paragraph("The vendor will use best efforts and provide reports monthly.")
    in_path = tmp_path / "c.docx"
    doc.save(in_path)

    main = json.dumps({"edits": [
        {"para": 0, "type": "replace", "find": "best efforts",
         "replace": "commercially reasonable efforts", "severity": "high",
         "comment": "main pass"},
    ]})
    completeness = json.dumps({"edits": [
        {"para": 0, "type": "replace", "find": "best efforts",  # duplicate -> dropped
         "replace": "commercially reasonable efforts", "severity": "high",
         "comment": "dup"},
        {"para": 0, "type": "replace", "find": "monthly",       # new -> added
         "replace": "quarterly", "severity": "medium", "comment": "missed by main"},
    ]})
    client = _StubClient([main, completeness])

    _, _, result = review(str(in_path), out_dir=str(tmp_path), client=client)
    assert len(client.calls) == 2
    # 2 unique edits (best efforts + monthly), both applied; no flags.
    assert result.num_applied == 2
    assert result.num_flagged == 0


def test_completeness_pass_can_be_disabled(tmp_path, monkeypatch):
    monkeypatch.setattr("review.COMPLETENESS_PASS", False)
    doc = Document()
    doc.add_paragraph("The term is best efforts to deliver.")
    in_path = tmp_path / "c.docx"
    doc.save(in_path)
    client = _StubClient(_response_two_edits())
    review(str(in_path), out_dir=str(tmp_path), client=client)
    assert len(client.calls) == 1  # only the main pass


def test_review_out_dir_is_respected(tmp_path):
    doc = Document()
    doc.add_paragraph("The term is best efforts to deliver.")
    in_path = tmp_path / "contract.docx"
    doc.save(in_path)
    out_dir = tmp_path / "outputs"

    client = _StubClient(_response_two_edits())
    out_docx, out_json, _ = review(str(in_path), out_dir=str(out_dir), client=client)
    assert out_docx.parent == out_dir
    assert out_json.parent == out_dir


def test_review_bytes_is_filesystem_free(tmp_path):
    """The server entry point: bytes in, bytes out, no paths."""
    doc = Document()
    doc.add_paragraph("The term is best efforts to deliver.")
    buf = __import__("io").BytesIO()
    doc.save(buf)
    in_bytes = buf.getvalue()

    client = _StubClient(_response_two_edits())
    result = review_bytes(in_bytes, client=client)

    assert isinstance(result.redlined_bytes, bytes)
    assert result.num_applied == 1
    assert result.num_flagged == 1
    # The returned bytes are a valid docx that accepts to the edited text.
    out = tmp_path / "from_bytes.docx"
    out.write_bytes(result.redlined_bytes)
    import subprocess
    accepted = subprocess.run(
        ["pandoc", "--track-changes=accept", "-t", "plain", str(out)],
        capture_output=True, text=True, check=True,
    ).stdout
    assert "commercially reasonable efforts" in accepted
