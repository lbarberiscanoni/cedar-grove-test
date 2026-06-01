"""Definition-of-done tests for the redline engine.

Each test builds its fixture contract in-process with python-docx so the assertions
target a known input. Validity (file opens) is a necessary check but NOT sufficient
— tests #2 and #4 are what prove the apply step is semantically correct.
"""

from __future__ import annotations

import os
import shutil
import subprocess
import sys
import zipfile
from pathlib import Path

import pytest
from docx import Document
from lxml import etree

# Make redline_engine importable when pytest runs from the repo root.
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from redline_engine import (  # noqa: E402
    Edit,
    apply_edits,
    extract_paragraphs,
    find_unique_match,
    normalize_with_map,
)

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _q(tag: str) -> str:
    return f"{{{W}}}{tag}"


# --- Fixtures -----------------------------------------------------------------


def _doc_bold_normal(tmp_path: Path) -> Path:
    """Paragraph: 'The term is best efforts to deliver.' where 'best' is bold."""
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("The term is ")
    r2 = p.add_run("best")
    r2.bold = True
    p.add_run(" efforts to deliver.")
    path = tmp_path / "bold.docx"
    doc.save(path)
    return path


def _doc_two_edits(tmp_path: Path) -> Path:
    """Single-run paragraph for the two-edits-in-one-paragraph regression test."""
    doc = Document()
    doc.add_paragraph("The term is 30 days and the fee is $500.")
    path = tmp_path / "two.docx"
    doc.save(path)
    return path


def _doc_duplicate(tmp_path: Path) -> Path:
    """Paragraph with a phrase that appears twice — must be flagged, not applied."""
    doc = Document()
    doc.add_paragraph("Either party may terminate. The party may also amend.")
    path = tmp_path / "dup.docx"
    doc.save(path)
    return path


def _doc_in_table(tmp_path: Path) -> Path:
    """Body + table cell, to confirm extraction sees paragraphs inside tables."""
    doc = Document()
    doc.add_paragraph("Pre-table paragraph.")
    t = doc.add_table(rows=1, cols=1)
    t.rows[0].cells[0].text = "The cell contains best efforts language."
    doc.add_paragraph("Post-table paragraph.")
    path = tmp_path / "table.docx"
    doc.save(path)
    return path


def _save_and_pandoc(tmp_path: Path, doc, mode: str, name: str = "out.docx") -> str:
    out = tmp_path / name
    doc.save(out)
    result = subprocess.run(
        ["pandoc", f"--track-changes={mode}", "-t", "plain", str(out)],
        capture_output=True, text=True, check=True,
    )
    return result.stdout


# --- Tests --------------------------------------------------------------------


def test_normalize_smart_quotes_and_double_spaces():
    s = "He said “hello”  with  spaces."
    n, _ = normalize_with_map(s)
    assert n == 'He said "hello" with spaces.'


def test_find_unique_match_returns_original_offsets():
    haystack = "best  efforts"  # two spaces — collapsed in normalized
    start, end = find_unique_match(haystack, "best efforts")
    assert haystack[start:end] == "best  efforts"


def test_find_unique_match_none_when_duplicate():
    assert find_unique_match("foo bar foo", "foo") is None


def test_bold_boundary_edit_preserves_formatting(tmp_path):
    """DoD #1: an edit whose `find` spans a bold/normal boundary applies correctly
    and preserves bold on the replacement.
    """
    in_path = _doc_bold_normal(tmp_path)
    doc = Document(in_path)
    paragraphs = extract_paragraphs(doc)
    assert paragraphs[0].text == "The term is best efforts to deliver."

    edits = [Edit(para=0, type="replace",
                  find="best efforts",
                  replace="commercially reasonable efforts")]
    flagged = apply_edits(doc, paragraphs, edits)
    assert flagged == []

    accepted = _save_and_pandoc(tmp_path, doc, "accept")
    rejected = _save_and_pandoc(tmp_path, doc, "reject")
    assert "The term is commercially reasonable efforts to deliver." in accepted
    assert "The term is best efforts to deliver." in rejected

    # Inspect XML: the <w:ins> should carry bold formatting (from the first
    # overlapped run, which was bold).
    out_path = tmp_path / "out.docx"
    with zipfile.ZipFile(out_path) as z:
        xml = z.read("word/document.xml")
    root = etree.fromstring(xml)
    ins_list = root.findall(f".//{_q('ins')}")
    assert len(ins_list) == 1, "exactly one insertion expected"
    rPr = ins_list[0].find(f".//{_q('rPr')}")
    assert rPr is not None, "insertion should carry rPr"
    assert rPr.find(_q("b")) is not None, "insertion rPr should preserve bold"

    # The <w:del> should contain TWO runs (one per crossed run-formatting boundary):
    # the bold "best" portion and the normal " efforts" portion.
    del_list = root.findall(f".//{_q('del')}")
    assert len(del_list) == 1
    del_runs = del_list[0].findall(_q("r"))
    assert len(del_runs) == 2, "deletion should split at formatting boundary"
    # First run keeps bold; second is normal.
    assert del_runs[0].find(f".//{_q('b')}") is not None
    assert del_runs[1].find(f".//{_q('rPr')}/{_q('b')}") is None


def test_two_edits_in_same_paragraph_apply_in_order(tmp_path):
    """DoD #2: the regression. Two edits in the same paragraph must both land in
    their correct positions — naive sequential application gets the second wrong.
    """
    in_path = _doc_two_edits(tmp_path)
    doc = Document(in_path)
    paragraphs = extract_paragraphs(doc)

    edits = [
        Edit(para=0, type="replace", find="30 days", replace="60 days"),
        Edit(para=0, type="replace", find="$500", replace="$1,000"),
    ]
    flagged = apply_edits(doc, paragraphs, edits)
    assert flagged == []

    accepted = _save_and_pandoc(tmp_path, doc, "accept")
    rejected = _save_and_pandoc(tmp_path, doc, "reject")
    # Assert the actual text order, not just that the file opens.
    assert "The term is 60 days and the fee is $1,000." in accepted, accepted
    assert "The term is 30 days and the fee is $500." in rejected, rejected


def test_two_edits_in_reverse_order_input_still_apply_correctly(tmp_path):
    """Same as above but with edits supplied in reverse document order — the engine
    must sort by span position internally so the result is identical.
    """
    in_path = _doc_two_edits(tmp_path)
    doc = Document(in_path)
    paragraphs = extract_paragraphs(doc)
    edits = [
        Edit(para=0, type="replace", find="$500", replace="$1,000"),
        Edit(para=0, type="replace", find="30 days", replace="60 days"),
    ]
    flagged = apply_edits(doc, paragraphs, edits)
    assert flagged == []
    accepted = _save_and_pandoc(tmp_path, doc, "accept")
    assert "The term is 60 days and the fee is $1,000." in accepted, accepted


def test_duplicate_find_is_flagged_not_applied(tmp_path):
    """DoD #3a: a `find` with 2 matches is flagged."""
    in_path = _doc_duplicate(tmp_path)
    doc = Document(in_path)
    paragraphs = extract_paragraphs(doc)

    edits = [Edit(para=0, type="replace", find="party", replace="Party")]
    flagged = apply_edits(doc, paragraphs, edits)
    assert len(flagged) == 1
    assert "2 matches" in flagged[0].reason

    # And the paragraph text should be unchanged.
    paragraphs_after = extract_paragraphs(doc)
    assert paragraphs_after[0].text == "Either party may terminate. The party may also amend."


def test_out_of_range_paragraph_is_flagged_not_applied(tmp_path):
    """DoD #3b: an out-of-range paragraph index is flagged."""
    in_path = _doc_two_edits(tmp_path)
    doc = Document(in_path)
    paragraphs = extract_paragraphs(doc)

    edits = [Edit(para=999, type="replace", find="anything", replace="something")]
    flagged = apply_edits(doc, paragraphs, edits)
    assert len(flagged) == 1
    assert "out of range" in flagged[0].reason


def test_delete_type_removes_text(tmp_path):
    """`type:"delete"` (or empty replace) emits only <w:del>, no <w:ins>."""
    doc = Document()
    doc.add_paragraph("Keep this. Drop the bracketed bit. And this stays.")
    in_path = tmp_path / "del.docx"
    doc.save(in_path)

    doc = Document(in_path)
    paragraphs = extract_paragraphs(doc)
    edits = [Edit(para=0, type="delete", find=" Drop the bracketed bit.")]
    flagged = apply_edits(doc, paragraphs, edits)
    assert flagged == []

    accepted = _save_and_pandoc(tmp_path, doc, "accept", name="del-out.docx")
    rejected = _save_and_pandoc(tmp_path, doc, "reject", name="del-out.docx")
    assert "Keep this. And this stays." in accepted
    assert "Keep this. Drop the bracketed bit. And this stays." in rejected

    out_path = tmp_path / "del-out.docx"
    with zipfile.ZipFile(out_path) as z:
        root = etree.fromstring(z.read("word/document.xml"))
    assert len(root.findall(f".//{_q('del')}")) == 1
    assert root.findall(f".//{_q('ins')}") == []


def test_table_paragraphs_are_indexed(tmp_path):
    """Extraction must include paragraphs inside table cells, in document order."""
    in_path = _doc_in_table(tmp_path)
    doc = Document(in_path)
    paragraphs = extract_paragraphs(doc)
    texts = [p.text for p in paragraphs]
    assert "Pre-table paragraph." in texts
    assert "The cell contains best efforts language." in texts
    assert "Post-table paragraph." in texts
    # Order: pre, cell, post
    pre = texts.index("Pre-table paragraph.")
    cell = texts.index("The cell contains best efforts language.")
    post = texts.index("Post-table paragraph.")
    assert pre < cell < post

    # And we can apply an edit to a paragraph inside the table cell.
    edits = [Edit(para=cell, type="replace",
                  find="best efforts",
                  replace="commercially reasonable efforts")]
    flagged = apply_edits(doc, paragraphs, edits)
    assert flagged == []
    accepted = _save_and_pandoc(tmp_path, doc, "accept", name="table-out.docx")
    assert "The cell contains commercially reasonable efforts language." in accepted


def test_round_trip_validity(tmp_path):
    """Saved docx is a valid OOXML file (would raise on malformed XML)."""
    in_path = _doc_bold_normal(tmp_path)
    doc = Document(in_path)
    paragraphs = extract_paragraphs(doc)
    edits = [Edit(para=0, type="replace",
                  find="best efforts",
                  replace="commercially reasonable efforts")]
    apply_edits(doc, paragraphs, edits)
    out = tmp_path / "rt.docx"
    doc.save(out)

    with zipfile.ZipFile(out) as z:
        xml = z.read("word/document.xml")
    etree.fromstring(xml)  # would raise if malformed
    Document(out)  # reopen via python-docx


@pytest.mark.skipif(shutil.which("soffice") is None, reason="LibreOffice (soffice) not installed")
def test_soffice_pdf_conversion(tmp_path):
    """DoD #5: soffice convert-to pdf succeeds (catches malformed XML in production)."""
    in_path = _doc_bold_normal(tmp_path)
    doc = Document(in_path)
    paragraphs = extract_paragraphs(doc)
    edits = [Edit(para=0, type="replace",
                  find="best efforts",
                  replace="commercially reasonable efforts")]
    apply_edits(doc, paragraphs, edits)
    out = tmp_path / "pdf-in.docx"
    doc.save(out)

    result = subprocess.run(
        ["soffice", "--headless", "--convert-to", "pdf", "--outdir", str(tmp_path), str(out)],
        capture_output=True, text=True, timeout=60,
    )
    assert result.returncode == 0, result.stderr
    pdf = tmp_path / out.with_suffix(".pdf").name
    assert pdf.exists()


def test_unicode_quote_in_doc_matches_ascii_in_find(tmp_path):
    """Model returns ASCII quotes in `find`; doc has smart quotes — normalize matches."""
    doc = Document()
    doc.add_paragraph("The clause says “best efforts” in section 3.")
    in_path = tmp_path / "smart.docx"
    doc.save(in_path)

    doc = Document(in_path)
    paragraphs = extract_paragraphs(doc)
    edits = [Edit(para=0, type="replace",
                  find='"best efforts"',  # ASCII quotes
                  replace='"commercially reasonable efforts"')]
    flagged = apply_edits(doc, paragraphs, edits)
    assert flagged == [], flagged

    accepted = _save_and_pandoc(tmp_path, doc, "accept", name="smart-out.docx")
    # After accept, the ASCII quotes from the replacement are present.
    assert '"commercially reasonable efforts"' in accepted
